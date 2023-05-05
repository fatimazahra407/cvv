from docx import Document
from docx.shared import Inches
document=Document()
document.add_picture('id.jpg' , width = Inches(1.5) )
name = 'Fatimat Ul Zahra'
phone_number = '+92 3365451035'
email = 'fatimatzahraa407@gmail.com'
document.add_paragraph (name).bold=True
document.add_paragraph ( phone_number + ' | ' + email )
document.add_heading ('About me')
document.add_paragraph  ('I am an Electrical engineer with the majors in computer engineering.I have strong foundation in math and programming skills. I made this profile to show my basic coding skills.  Currently I am looking for position in lab as a researcher in the feild of Artificial Intellegence and Data Science.')
document.add_heading('Education')
p = document.add_paragraph ()
Degree_name = 'BS In Electrical Engineering '
Grade = 'Grade A' 
School_name ='International Islamic University Islamabad'
Duration = '  2018~ 2022'
p.add_run(Degree_name + ' ' + '\n').bold=True
p.add_run(Grade + ' '  +'\n')
p.add_run(School_name + '  ' +'\t'  )
p.add_run(Duration+ '  ' ).italic = True
document.add_heading('projects')
q=document.add_paragraph()
q.add_run('Real Time Face emotion recognition').bold=True
q.add_run('  Sep 2022\n').italic = True
q.add_run (' --> Deep Learning \n'
+ ' --> DeepFace \n'
 + ' --> Opencv \n'
)
q=document.add_paragraph()
q.add_run('Movie Recommendation System').bold=True
q.add_run('  Aug 2022 \n').italic = True
q.add_run (' --> Clustering models \n'
+ ' --> K-nearest neighbors \n'
 + ' --> Bayesian networks\n '
)
q=document.add_paragraph()
q.add_run('Predictive Intelligence-based Store Management System').bold=True
q.add_run('  January 2023 \n').italic = True
q.add_run (' --> Image Recognition \n'
+ ' --> Recommender systems \n'
 + ' --> Deep Learning\n ' +
 ' --> Edge Computation \n '
)


document.save('cv.docx')
